"""Document format detection and preprocessing."""
import mimetypes
from pathlib import Path
from typing import Tuple, Optional
from PIL import Image
from docx import Document as DocxDocument
import io
import concurrent.futures
try:
    import pypdfium2 as pdfium
    PDFIUM_AVAILABLE = True
except ImportError:
    PDFIUM_AVAILABLE = False


class DocumentProcessor:
    """Handles document format detection and preprocessing."""
    
    SUPPORTED_FORMATS = {
        'image/png': 'PNG',
        'image/jpeg': 'JPEG',
        'image/jpg': 'JPEG',
        'application/pdf': 'PDF',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'DOCX',
        'image/tiff': 'TIFF',
        'image/tif': 'TIFF'
    }
    
    def detect_format(self, file_path: str) -> Tuple[str, Optional[str]]:
        """Detect document format and return (format_type, mime_type)."""
        path = Path(file_path)
        mime_type, _ = mimetypes.guess_type(str(path))
        
        if mime_type in self.SUPPORTED_FORMATS:
            return self.SUPPORTED_FORMATS[mime_type], mime_type
        
        # Fallback: check extension
        ext = path.suffix.lower()
        ext_map = {
            '.png': ('PNG', 'image/png'),
            '.jpg': ('JPEG', 'image/jpeg'),
            '.jpeg': ('JPEG', 'image/jpeg'),
            '.pdf': ('PDF', 'application/pdf'),
            '.docx': ('DOCX', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            '.tiff': ('TIFF', 'image/tiff'),
            '.tif': ('TIFF', 'image/tiff')
        }
        
        if ext in ext_map:
            return ext_map[ext]
        
        raise ValueError(f"Unsupported file format: {mime_type or ext}")
    
    def preprocess_for_ocr(self, file_path: str, file_type: str) -> list:
        """Preprocess document for OCR processing.
        
        Returns list of file contents (as bytes) for OCR processing.
        For PDFs, returns the entire PDF file (OpenAI can process PDFs directly).
        For images, returns image bytes.
        For DOCX, returns extracted text.
        """
        if file_type in ['PNG', 'JPEG', 'TIFF']:
            # For images, return the image file
            with open(file_path, 'rb') as f:
                return [f.read()]
        
        elif file_type == 'PDF':
            # Convert PDF pages to images for OpenAI Vision API (parallel processing)
            # Note: OpenAI Vision API doesn't support PDFs directly via image_url
            # We need to convert PDF pages to images
            if not PDFIUM_AVAILABLE:
                raise ValueError("pypdfium2 is required for PDF processing. Install with: pip install pypdfium2")

            from app.utils.logger import get_logger
            logger = get_logger(__name__)

            pdf = None
            try:
                # Open PDF with error handling
                logger.info(f"Opening PDF: {file_path}")
                pdf = pdfium.PdfDocument(file_path)
                page_count = len(pdf)
                logger.info(f"PDF has {page_count} pages")

                # Limit maximum pages to prevent memory issues
                MAX_PAGES = 50
                if page_count > MAX_PAGES:
                    logger.warning(f"PDF has {page_count} pages, limiting to first {MAX_PAGES} pages")
                    page_count = MAX_PAGES

                def convert_page_to_image(page_num):
                    """Convert a single PDF page to image bytes with error handling."""
                    page = None
                    bitmap = None
                    pil_image = None

                    try:
                        logger.debug(f"Converting page {page_num + 1}/{page_count}")
                        page = pdf.get_page(page_num)

                        # Try rendering at 150 DPI first (safer for complex graphics)
                        # If this fails, we'll catch and skip the page
                        try:
                            # Use lower DPI for initial attempt (150 instead of 200)
                            # This is safer for PDFs with complex graphics/charts
                            bitmap = page.render(scale=150/72)  # 72 is default DPI
                        except Exception as render_error:
                            logger.warning(f"Page {page_num + 1}: Render at 150 DPI failed, trying 100 DPI: {render_error}")
                            # Fallback to even lower DPI for very complex pages
                            try:
                                bitmap = page.render(scale=100/72)
                            except Exception as fallback_error:
                                raise Exception(f"Failed to render at both 150 and 100 DPI: {fallback_error}")

                        # Convert bitmap to PIL image
                        pil_image = bitmap.to_pil()

                        # Limit image size to prevent memory issues
                        # If image is too large, resize it
                        max_dimension = 4000  # Max width or height in pixels
                        width, height = pil_image.size
                        if width > max_dimension or height > max_dimension:
                            logger.info(f"Page {page_num + 1}: Resizing from {width}x{height} to fit {max_dimension}px")
                            if width > height:
                                new_width = max_dimension
                                new_height = int(height * (max_dimension / width))
                            else:
                                new_height = max_dimension
                                new_width = int(width * (max_dimension / height))
                            pil_image = pil_image.resize((new_width, new_height), Image.Resampling.LANCZOS)

                        # Convert to bytes with optimization
                        img_bytes = io.BytesIO()
                        # Use JPEG for complex graphics (smaller, faster)
                        # Quality 85 is good balance between size and OCR accuracy
                        pil_image.save(img_bytes, format='JPEG', quality=85, optimize=True)

                        logger.debug(f"Page {page_num + 1} converted successfully ({len(img_bytes.getvalue())} bytes)")
                        return page_num, img_bytes.getvalue(), None

                    except Exception as e:
                        error_msg = f"Error converting page {page_num + 1}: {str(e)}"
                        logger.error(error_msg)
                        return page_num, None, error_msg
                    finally:
                        # Clean up resources in reverse order
                        if pil_image:
                            try:
                                pil_image.close()
                            except:
                                pass
                        if bitmap:
                            try:
                                del bitmap  # Free bitmap memory
                            except:
                                pass
                        if page:
                            try:
                                page.close()
                            except:
                                pass

                # Convert pages sequentially first (safer than parallel for problematic PDFs)
                # Use sequential processing to avoid segfaults from parallel C++ library calls
                images = [None] * page_count
                errors = []

                logger.info(f"Converting {page_count} pages sequentially (safer for complex PDFs)")
                for i in range(page_count):
                    page_num, image_bytes, error = convert_page_to_image(i)
                    if error:
                        errors.append(f"Page {page_num + 1}: {error}")
                        # Create a placeholder for failed pages
                        images[page_num] = None
                    else:
                        images[page_num] = image_bytes

                # Filter out None values (failed pages)
                images = [img for img in images if img is not None]

                if len(images) == 0:
                    raise Exception(f"Failed to convert any PDF pages. Errors: {'; '.join(errors[:3])}")

                if errors:
                    logger.warning(f"Successfully converted {len(images)}/{page_count} pages. Some pages failed: {errors[:3]}")
                else:
                    logger.info(f"Successfully converted all {len(images)} pages")

                return images

            except Exception as e:
                error_msg = f"PDF processing failed: {str(e)}"
                logger.error(error_msg)
                raise Exception(error_msg)
            finally:
                # Always close the PDF to free resources
                if pdf:
                    try:
                        pdf.close()
                        logger.debug("PDF document closed")
                    except Exception as e:
                        logger.warning(f"Error closing PDF: {e}")
        
        elif file_type == 'DOCX':
            # For DOCX, extract text directly (OpenAI can process text)
            doc = DocxDocument(file_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])
            return [text_content.encode('utf-8')]
        
        else:
            raise ValueError(f"Unsupported file type for preprocessing: {file_type}")
    
    def get_file_size(self, file_path: str) -> int:
        """Get file size in bytes."""
        return Path(file_path).stat().st_size

